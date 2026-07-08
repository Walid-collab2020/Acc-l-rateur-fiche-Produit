"""
Extracteur de texte universel — aucune perte d'information, tous types de documents.

Stratégie par type :
- PDF      → extraction page par page avec marqueurs --- PAGE N ---
- DOCX     → paragraphes avec marqueurs === SECTION === sur les titres (styles Word)
- DOC      → win32com avec détection de sections par OutlineLevel + heuristiques
- XLSX/XLS → si fiche paramétrage BOSS → extracteur dédié (voir excel_parametrage_extractor)
             sinon → feuille par feuille avec en-têtes de colonnes comme contexte
- TXT      → lecture directe
"""
import io
import logging
import re
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text(file_path: str, mime_type: str) -> tuple[str, int]:
    """Extract text from a document. Returns (text, page_count)."""
    path = Path(file_path)
    if not path.exists():
        return "", 0

    try:
        ext = path.suffix.lower()
        if mime_type == "application/pdf" or ext == ".pdf":
            return _extract_pdf(file_path)
        elif mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ) or ext == ".docx":
            return _extract_docx(file_path)
        elif mime_type == "application/msword" or ext == ".doc":
            return _extract_doc(file_path)
        elif mime_type in (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/vnd.ms-excel",
        ) or ext in (".xlsx", ".xls"):
            return _extract_excel(file_path)
        elif ext == ".txt":
            text = path.read_text(encoding="utf-8", errors="ignore")
            return text, 1
        else:
            return "", 0
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {e}")
        return "", 0


# ─── PDF ──────────────────────────────────────────────────────────────────────

def _extract_pdf(file_path: str) -> tuple[str, int]:
    """Page-by-page extraction with --- PAGE N --- markers."""
    # pdfminer page-by-page (preferred — pure Python)
    try:
        from pdfminer.pdfpage import PDFPage
        from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
        from pdfminer.converter import TextConverter
        from pdfminer.layout import LAParams

        rsrcmgr = PDFResourceManager()
        laparams = LAParams()
        pages_text = []

        with open(file_path, "rb") as f:
            for page_num, page in enumerate(PDFPage.get_pages(f), start=1):
                output = io.StringIO()
                device = TextConverter(rsrcmgr, output, laparams=laparams)
                interpreter = PDFPageInterpreter(rsrcmgr, device)
                interpreter.process_page(page)
                device.close()
                page_text = output.getvalue().strip()
                if page_text:
                    pages_text.append(f"\n\n--- PAGE {page_num} ---\n{page_text}")

        if pages_text:
            return "".join(pages_text), len(pages_text)
    except ImportError:
        pass
    except Exception as e:
        logger.error(f"pdfminer page-by-page error on {file_path}: {e}")

    # Fallback: PyMuPDF with page markers
    try:
        import fitz
        doc = fitz.open(file_path)
        parts = []
        for i, page in enumerate(doc):
            t = page.get_text().strip()
            if t:
                parts.append(f"\n\n--- PAGE {i+1} ---\n{t}")
        doc.close()
        return "".join(parts), len(doc)
    except Exception as e:
        logger.error(f"PyMuPDF error on {file_path}: {e}")

    return "", 0


# ─── DOCX ─────────────────────────────────────────────────────────────────────

def _extract_docx(file_path: str) -> tuple[str, int]:
    """
    Extract DOCX preserving structure:
    - Titres (styles Heading/Titre) → marqueurs === SECTION ===
    - Corps de texte → paragraphes
    - Tableaux → ligne par ligne avec en-têtes si disponibles
    """
    try:
        from docx import Document
        doc = Document(file_path)
        parts = []

        current_page = 1
        parts.append(f"\n\n--- PAGE {current_page} ---\n")

        for para in doc.paragraphs:
            # Détecter les sauts de page explicites dans les runs du paragraphe
            for run in para.runs:
                try:
                    from docx.oxml.ns import qn as _qn
                    for br in run._element.findall(_qn("w:br")):
                        if br.get(_qn("w:type")) == "page":
                            current_page += 1
                            parts.append(f"\n\n--- PAGE {current_page} ---\n")
                except Exception:
                    pass

            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name or "").lower() if para.style else ""
            # Détecter les titres par style Word
            is_heading = (
                "heading" in style
                or "titre" in style
                or "title" in style
                or style.startswith("h1") or style.startswith("h2")
            )
            # Heuristique : ligne courte en MAJUSCULES = titre probable
            is_allcaps_title = (
                text.isupper() and len(text) < 120
                and not text.startswith("–") and not text.startswith("-")
            )
            # Heuristique : commence par "ARTICLE", "PARTIE", "SECTION", "I.", "II.", "1.", "§"
            is_numbered_title = bool(re.match(
                r"^(ARTICLE\s+\d|PARTIE\s+[IVX\d]|SECTION\s+[IVX\d]|§\s*\d|[IVX]+\s*[-\.]\s+[A-Z]|\d+\s*\.\s+[A-Z])",
                text
            ))

            if is_heading or is_allcaps_title or is_numbered_title:
                parts.append(f"\n\n=== SECTION : {text} ===\n")
            else:
                parts.append(text + "\n")

        # Tableaux avec en-têtes
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    rows.append(cells)
            if not rows:
                continue
            # Première ligne comme en-tête si elle ressemble à des libellés
            header = rows[0]
            if len(rows) > 1:
                parts.append(f"\n[Tableau {table_idx+1} — colonnes : {' | '.join(header)}]\n")
                for row in rows[1:]:
                    line = " | ".join(
                        f"{h}: {v}" for h, v in zip(header, row) if v
                    )
                    if line:
                        parts.append(line + "\n")
            else:
                line = " | ".join(c for c in header if c)
                if line:
                    parts.append(line + "\n")

        full_text = "".join(parts)

        # Essai 1 : propriétés intégrées du document (Pages)
        page_count = 0
        try:
            cp = doc.core_properties
            # python-docx expose revision, author, etc. mais pas les pages nativement
            # On tente via les custom properties ou l'XML des statistiques
            from docx.oxml.ns import qn
            app_xml = doc.part.package.part_related_by(
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties"
            )
            pages_el = app_xml._element.find(qn("Pages"))
            if pages_el is not None and pages_el.text:
                page_count = int(pages_el.text)
        except Exception:
            pass

        # Essai 2 : estimation par taille de texte (≈ 2 500 chars / page A4)
        if not page_count:
            page_count = max(1, len(full_text) // 2500)

        return full_text, page_count

    except ImportError:
        logger.warning("python-docx not installed")
        return "", 0
    except Exception as e:
        logger.error(f"python-docx extraction failed for {file_path}: {e}")
        return "", 0


# ─── DOC (binaire Word legacy) ────────────────────────────────────────────────

def _extract_doc(file_path: str) -> tuple[str, int]:
    """
    Extract .doc via win32com en préservant la structure :
    - OutlineLevel 1-3 → marqueurs === SECTION ===
    - Tableaux → lignes avec en-têtes
    Fallback : texte brut avec heuristiques de détection de titres.
    """
    path = Path(file_path)

    # Essai 1 : win32com avec structure (OutlineLevel + tableaux)
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(path.resolve()))

        parts = []
        current_page = 0  # track page changes to inject --- PAGE N --- markers

        # Paragraphes avec niveaux de titre
        for para in doc.Paragraphs:
            text = para.Range.Text.strip().replace("\r", "").replace("\x07", "")

            # wdActiveEndPageNumber = 3 — page réelle du paragraphe dans Word
            try:
                para_page = int(para.Range.Information(3))
            except Exception:
                para_page = current_page

            if para_page != current_page:
                current_page = para_page
                parts.append(f"\n\n--- PAGE {current_page} ---\n")

            if not text:
                continue

            try:
                outline_level = para.OutlineLevel  # 1-9 = titre, 10 = corps
            except Exception:
                outline_level = 10

            style_name = ""
            try:
                style_name = para.Style.NameLocal.lower()
            except Exception:
                pass

            is_heading = (
                outline_level <= 3
                or "titre" in style_name
                or "heading" in style_name
            )
            # Heuristique de secours si OutlineLevel = 10 mais ressemble à un titre
            if not is_heading:
                is_heading = bool(re.match(
                    r"^(ARTICLE\s+\d|PARTIE\s+[IVX\d]|§\s*\d|[IVX]+\s*[-\.]\s+[A-Z]|\d+\s*\.\s+[A-Z])",
                    text
                )) or (text.isupper() and len(text) < 100)

            if is_heading:
                parts.append(f"\n\n=== SECTION : {text} ===\n")
            else:
                parts.append(text + "\n")

        # Tableaux
        for t_idx in range(1, doc.Tables.Count + 1):
            try:
                table = doc.Tables(t_idx)
                rows = []
                for r in range(1, table.Rows.Count + 1):
                    row_cells = []
                    for c in range(1, table.Columns.Count + 1):
                        try:
                            cell_text = table.Cell(r, c).Range.Text.strip().replace("\r", "").replace("\x07", "")
                            row_cells.append(cell_text)
                        except Exception:
                            row_cells.append("")
                    if any(row_cells):
                        rows.append(row_cells)
                if rows:
                    header = rows[0]
                    parts.append(f"\n[Tableau {t_idx} — {' | '.join(header)}]\n")
                    for row in rows[1:]:
                        line = " | ".join(f"{h}: {v}" for h, v in zip(header, row) if v)
                        if line:
                            parts.append(line + "\n")
            except Exception:
                pass

        # wdStatisticPages = 2 — nombre de pages calculé par Word
        try:
            page_count = doc.ComputeStatistics(2)
        except Exception:
            page_count = max(1, len("".join(parts)) // 2500)

        doc.Close(False)
        word.Quit()
        pythoncom.CoUninitialize()

        full_text = "".join(parts)
        return full_text, page_count

    except ImportError:
        logger.warning("win32com not available — trying fallback for .doc")
    except Exception as e:
        logger.error(f"win32com .doc extraction failed for {file_path}: {e}")

    # Fallback : antiword ou extraction brute avec heuristiques de titres
    try:
        import subprocess
        result = subprocess.run(
            ["antiword", str(path)],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0 and result.stdout:
            return _inject_section_markers(result.stdout), 1
    except Exception:
        pass

    return "", 0


# ─── EXCEL ────────────────────────────────────────────────────────────────────

def _cell_value_to_str(cell) -> str:
    """
    Convert an openpyxl Cell to a human-readable string.
    Cells formatted as percentage (number_format contains '%') are converted:
      0.03 with format '0%'  → '3 %'
      0.045 with format '0.00%' → '4,50 %'
    Other numeric values are stringified as-is.
    """
    v = cell.value
    if v is None:
        return ""
    if isinstance(v, (int, float)) and not isinstance(v, bool):
        fmt = getattr(cell, "number_format", None) or ""
        if "%" in fmt:
            pct = v * 100
            if pct == int(pct):
                return f"{int(pct)} %"
            else:
                formatted = f"{pct:.2f}".replace(".", ",")
                return f"{formatted} %"
    return str(v).strip()


def _extract_excel(file_path: str) -> tuple[str, int]:
    """
    Extract Excel :
    - Si structure paramétrage BOSS détectée → signal pour extracteur dédié
      (retourne texte brut minimal, l'extracteur dédié est appelé en amont dans referentiel_service)
    - Sinon → feuille par feuille avec en-têtes de colonnes comme contexte de chaque ligne
    Amélioration : les cellules formatées en pourcentage (0.03 = 3 %) sont converties correctement.
    """
    try:
        import openpyxl
        # data_only=True reads cached values (not formulas)
        # NOT read_only so cell.number_format is accessible for percentage detection
        wb = openpyxl.load_workbook(file_path, data_only=True)
        sheets_text = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_data = []
            header_row: Optional[list[str]] = None
            header_row_idx = -1

            for row_idx, row in enumerate(ws.iter_rows()):
                cells = [_cell_value_to_str(c) for c in row]
                if not any(cells):
                    continue

                # Détecter la ligne d'en-tête (première ligne avec plusieurs cellules non vides)
                if header_row is None and sum(1 for c in cells if c) >= 2:
                    header_row = cells
                    header_row_idx = row_idx
                    rows_data.append(f"[En-têtes] : {' | '.join(c for c in cells if c)}")
                    continue

                # Pour les lignes de données : associer chaque valeur à son en-tête
                if header_row:
                    line_parts = []
                    for h, v in zip(header_row, cells):
                        if v:
                            label = h.strip() if h.strip() else f"Col{header_row.index(h)+1}"
                            line_parts.append(f"{label}: {v}")
                    if line_parts:
                        rows_data.append(" | ".join(line_parts))
                else:
                    # Pas encore d'en-tête trouvé
                    line = " | ".join(c for c in cells if c)
                    if line:
                        rows_data.append(line)

            if rows_data:
                sheets_text.append(
                    f"\n\n=== FEUILLE : {sheet_name} ===\n" + "\n".join(rows_data)
                )

        wb.close()
        return "\n".join(sheets_text), len(wb.sheetnames)

    except ImportError:
        logger.warning("openpyxl not installed")
        return f"[Excel non lisible — openpyxl manquant : {file_path}]", 0
    except Exception as e:
        logger.error(f"Excel extraction failed for {file_path}: {e}")
        return "", 0


# ─── Utilitaire : injection de marqueurs de section heuristiques ──────────────

def _inject_section_markers(text: str) -> str:
    """
    Pour les textes bruts sans structure (fallback .doc) :
    détecte les titres par heuristiques et injecte === SECTION === .
    """
    lines = text.split("\n")
    result = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            result.append("")
            continue
        is_title = (
            (stripped.isupper() and 4 < len(stripped) < 100)
            or bool(re.match(
                r"^(ARTICLE\s+\d|PARTIE\s+[IVX\d]|§\s*\d|[IVX]+\s*[-.]\s+[A-Z]|\d+\s*\.\s+[A-ZÀ-ÿ])",
                stripped
            ))
        )
        if is_title:
            result.append(f"\n=== SECTION : {stripped} ===")
        else:
            result.append(stripped)
    return "\n".join(result)
